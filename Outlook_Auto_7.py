import win32com.client
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.probability import FreqDist
import re
import sys
import traceback
import pythoncom
from datetime import datetime, timedelta
from docx import Document
import os
import ssl
ssl._create_default_https_context = ssl._create_unverified_context

def log(message):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {message}")

def download_nltk_data():
    try:
        log("Starting NLTK data download...")
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
        log("NLTK data download completed.")
    except Exception as e:
        log(f"Error downloading NLTK data: {str(e)}")
        sys.exit(1)

# Disable SSL verification (use with caution)
ssl._create_default_https_context = ssl._create_unverified_context

download_nltk_data()

def connect_to_outlook():
    try:
        log("Connecting to Outlook...")
        outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        log("Successfully connected to Outlook.")
        return outlook
    except Exception as e:
        log(f"Error connecting to Outlook: {str(e)}")
        log("Please ensure Outlook is installed and running.")
        sys.exit(1)

def fetch_emails(outlook, search_criteria, days_back=30):
    try:
        log(f"Fetching emails related to '{search_criteria}'...")
        folders_to_search = [
            (6, "Inbox"),
            (5, "Sent Items"),
            (3, "Deleted Items"),
            (4, "Outbox"),
            (2, "Drafts")
        ]
        
        all_messages = []
        start_date = (datetime.now() - timedelta(days=days_back)).strftime("%m/%d/%Y")
        
        for folder_const, folder_name in folders_to_search:
            try:
                log(f"Searching in {folder_name}...")
                folder = outlook.GetDefaultFolder(folder_const)
                messages = folder.Items
                messages.Sort("[ReceivedTime]", True)
                
                log(f"Applying filter for '{search_criteria}' in {folder_name}...")
                filter_string = (f"@SQL=((\"urn:schemas:httpmail:subject\" LIKE '%{search_criteria}%') OR "
                                 f"(\"urn:schemas:httpmail:textdescription\" LIKE '%{search_criteria}%') OR "
                                 f"(\"urn:schemas:httpmail:fromname\" LIKE '%{search_criteria}%') OR "
                                 f"(\"urn:schemas:httpmail:fromaddress\" LIKE '%{search_criteria}%')) AND "
                                 f"\"urn:schemas:httpmail:datereceived\" >= '{start_date}'")
                filtered_messages = messages.Restrict(filter_string)
                
                log(f"Found {filtered_messages.Count} emails in {folder_name}")
                all_messages.extend(list(filtered_messages))
            except Exception as e:
                log(f"Error searching {folder_name}: {str(e)}")
                continue
        
        log(f"Total emails found across all folders: {len(all_messages)}")
        return all_messages
    except Exception as e:
        log(f"Error fetching emails: {str(e)}")
        return []

def get_email_content(message):
    try:
        subject = message.Subject
        sender = message.SenderEmailAddress
        body = message.Body
        received_time = message.ReceivedTime
        return subject, sender, body, received_time
    except Exception as e:
        log(f"Error extracting email content: {str(e)}")
        return None, None, None, None

def preprocess_text(text):
    try:
        log("Preprocessing text...")
        text = re.sub(r'[^\w\s]', '', text.lower())
        tokens = word_tokenize(text)
        stop_words = set(stopwords.words('english'))
        tokens = [word for word in tokens if word not in stop_words]
        log("Text preprocessing completed.")
        return tokens
    except Exception as e:
        log(f"Error preprocessing text: {str(e)}")
        return []

def simple_summarize(text, num_sentences=3):
    sentences = sent_tokenize(text)
    word_frequencies = FreqDist(word_tokenize(text.lower()))
    sentence_scores = {}
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in word_frequencies:
                if sentence not in sentence_scores:
                    sentence_scores[sentence] = word_frequencies[word]
                else:
                    sentence_scores[sentence] += word_frequencies[word]
    summary_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:num_sentences]
    summary = ' '.join(summary_sentences)
    return summary

def extract_key_info(text, search_term):
    try:
        log("Extracting key information...")
        status_keywords = ['open', 'in progress', 'resolved', 'closed']
        status = next((word for word in status_keywords if word in text.lower()), "N/A")

        log("Generating summary...")
        summary = simple_summarize(text)
        log("Key information extraction completed.")

        return {
            'search_term': search_term,
            'status': status,
            'summary': summary
        }
    except Exception as e:
        log(f"Error extracting key info: {str(e)}")
        return {
            'search_term': search_term,
            'status': 'N/A',
            'summary': 'Error generating summary'
        }

def analyze_emails(search_term, days_back=30):
    try:
        log("Initializing COM library...")
        pythoncom.CoInitialize()
        outlook = connect_to_outlook()
        messages = fetch_emails(outlook, search_term, days_back)

        if not messages:
            log(f"No emails found for search term: {search_term}")
            return None, None

        log("Processing emails...")
        all_content = ""
        email_count = 0
        latest_date = None

        for message in messages:
            subject, sender, body, received_time = get_email_content(message)
            if body:
                all_content += f"Subject: {subject}\nFrom: {sender}\n{body}\n\n"
                email_count += 1
                if latest_date is None or received_time > latest_date:
                    latest_date = received_time

        tokens = preprocess_text(all_content)
        log("Analyzing word frequency...")
        freq_dist = FreqDist(tokens)
        common_words = freq_dist.most_common(10)

        log("Extracting key information...")
        key_info = extract_key_info(all_content, search_term)
        key_info['email_count'] = email_count
        key_info['latest_date'] = latest_date.strftime("%Y-%m-%d %H:%M:%S") if latest_date else "No emails found"

        log("Analysis completed.")
        return key_info, common_words
    except Exception as e:
        log(f"Error analyzing emails: {str(e)}")
        return None, None
    finally:
        log("Uninitializing COM library...")
        pythoncom.CoUninitialize()

def export_to_word(key_info, common_words, output_file):
    try:
        log("Exporting results to Word document...")
        doc = Document()
        doc.add_heading(f"Email Analysis: {key_info['search_term']}", 0)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(f"Search Term: {key_info['search_term']}")
        doc.add_paragraph(f"Number of related emails: {key_info['email_count']}")
        doc.add_paragraph(f"Latest email date: {key_info['latest_date']}")
        doc.add_paragraph(f"Status (if applicable): {key_info['status']}")
        doc.add_paragraph(f"Summary: {key_info['summary']}")

        doc.add_heading("Most Common Words", level=1)
        for word, count in common_words:
            doc.add_paragraph(f"{word}: {count}")

        doc.save(output_file)
        log(f"Results exported to {output_file}")
    except Exception as e:
        log(f"Error exporting to Word: {str(e)}")

def main():
    try:
        search_term = input("Enter the search term (incident number, keyword, etc.): ")
        days_back = int(input("Enter the number of days to search back (default is 30): ") or 30)
        log(f"Analyzing emails for search term: {search_term}")

        key_info, common_words = analyze_emails(search_term, days_back)

        if key_info and common_words:
            log("Printing analysis results...")
            print(f"\nSummary for '{search_term}':")
            print(f"Number of related emails: {key_info['email_count']}")
            print(f"Latest email date: {key_info['latest_date']}")
            print(f"Status (if applicable): {key_info['status']}")
            print(f"Summary: {key_info['summary']}")
            print("\nMost common words:")
            for word, count in common_words:
                print(f"{word}: {count}")

            output_file = f"Email_Analysis_{search_term.replace(' ', '_')}.docx"
            export_to_word(key_info, common_words, output_file)
        else:
            log("Unable to analyze emails. Please check the search term and try again.")
    except Exception as e:
        log(f"An unexpected error occurred: {str(e)}")
        log("Error details:")
        traceback.print_exc()

if __name__ == "__main__":
    main()